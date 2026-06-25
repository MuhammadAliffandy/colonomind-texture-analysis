import os
from docx import Document
from docx.shared import Inches

def main():
    base_dir = "/Users/aliffandy/Documents/PukulEnam/Colonomind-Texture-analysis"
    docx_path = os.path.join(base_dir, "reports", "Texture_Analysis_Medical_Report.docx")
    output_path = os.path.join(base_dir, "reports", "Texture_Analysis_Medical_Report_Updated.docx")
    figures_dir = os.path.join(base_dir, "Texture_Analysis_Presentation_Figures")

    if not os.path.exists(docx_path):
        print(f"Error: Could not find {docx_path}")
        return

    # Load existing document
    doc = Document(docx_path)
    
    # Add a page break and a new heading for the figures
    doc.add_page_break()
    doc.add_heading('Presentation Figures', level=1)
    
    # Get all png files and sort them
    figures = [f for f in os.listdir(figures_dir) if f.endswith('.png')]
    figures.sort()
    
    if not figures:
        print("No figures found in the directory.")
        return
        
    for fig in figures:
        # Add heading based on filename
        fig_title = fig.replace('.png', '').replace('_', ' ')
        doc.add_heading(fig_title, level=2)
        
        # Add the image
        img_path = os.path.join(figures_dir, fig)
        # Adding image with width of 6 inches to fit well on standard page
        doc.add_picture(img_path, width=Inches(6.0))
        # Add some space after the image
        doc.add_paragraph("")
        
    # Save the updated document
    doc.save(output_path)
    print(f"Successfully updated document and saved to {output_path}")

if __name__ == '__main__':
    main()
