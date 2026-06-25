import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def add_table_with_style(doc, title, headers, data):
    doc.add_heading(title, level=2)
    
    table = doc.add_table(rows=1, cols=len(headers))
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        # Optional: make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    doc.add_paragraph("") # empty line after table

def main():
    base_dir = "/Users/aliffandy/Documents/PukulEnam/Colonomind-Texture-analysis"
    docx_path = os.path.join(base_dir, "reports", "Texture_Analysis_Medical_Report_Updated.docx")
    output_path = os.path.join(base_dir, "reports", "Texture_Analysis_Medical_Report_Final.docx")

    if not os.path.exists(docx_path):
        print(f"Error: Could not find {docx_path}")
        return

    doc = Document(docx_path)
    
    doc.add_page_break()
    doc.add_heading('Rule-Based Thresholds', level=1)

    # LIMUC Data
    limuc_headers = ["MES Level", "HH_Ent Range (Avg)", "HL_Mean Range (Avg)", "LH_Mean Range (Avg)"]
    limuc_data = [
        ["MES 0", "8.2590 - 8.5445 (8.3974)", "0.0143 - 0.2994 (0.1681)", "-0.9094 - -0.6555 (-0.7798)"],
        ["MES 1", "8.3210 - 8.5969 (8.4470)", "0.1472 - 0.3485 (0.2244)", "-0.9532 - -0.7269 (-0.8306)"],
        ["MES 2", "8.4059 - 8.6429 (8.5059)", "0.0088 - 0.3517 (0.2015)", "-0.9605 - -0.6981 (-0.8180)"],
        ["MES 3", "8.4241 - 8.6480 (8.5257)", "0.1181 - 0.3610 (0.2154)", "-0.9493 - -0.7229 (-0.8304)"]
    ]
    
    add_table_with_style(doc, "LIMUC Rule-Based Thresholds", limuc_headers, limuc_data)

    # TMC Data
    tmc_headers = ["MES Level", "HL_Ent Range (Avg)", "LH_Ent Range (Avg)", "GLCM_Dissimilarity Range (Avg)"]
    tmc_data = [
        ["MES 0", "8.7361 - 8.9130 (8.8199)", "8.7465 - 8.9182 (8.8281)", "9.9851 - 14.7280 (12.5691)"],
        ["MES 1", "8.8140 - 8.9995 (8.8944)", "8.8204 - 9.0033 (8.8999)", "8.8356 - 13.7776 (11.6142)"],
        ["MES 2", "8.8641 - 9.0111 (8.9263)", "8.8710 - 9.0144 (8.9311)", "9.1047 - 13.4166 (11.4944)"],
        ["MES 3", "8.8707 - 8.9912 (8.9202)", "8.8735 - 8.9940 (8.9244)", "11.0037 - 17.0195 (14.3315)"]
    ]
    
    add_table_with_style(doc, "TMC Rule-Based Thresholds", tmc_headers, tmc_data)

    doc.save(output_path)
    print(f"Successfully updated document and saved to {output_path}")

if __name__ == '__main__':
    main()
